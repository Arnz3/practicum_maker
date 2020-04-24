from tkinter import *
from tkinter import filedialog
from tkinter.filedialog import askopenfilename
from docx import *
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path
import re

nr = 0
path = ""
doc = Document()
voorb_f = "niets"

# --------   styles   --------

# create style "Titel"
titel = doc.styles.add_style("Titel", WD_STYLE_TYPE.PARAGRAPH)
font = titel.font
font.name = "Verdana"
font.size = Pt(20)

# create style "Normal"
normal = doc.styles["Normal"]
font1 = normal.font
font1.name = "Verdana"
font1.size = Pt(12)

# create style "specs"
specs = doc.styles.add_style("specs", WD_STYLE_TYPE.PARAGRAPH)
font2 = specs.font
font2.name = "Verdana"
font2.size = Pt(14)


def add_titel(text):
    global doc
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.style = doc.styles["Titel"]
    p.alignment = 1
    run = p.add_run(text)
    run.bold = True
    run.underline = True


def add_spec(text):
    global doc
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.style = doc.styles["specs"]
    p.alignment = 1
    run = p.add_run(text)
    run.bold = True
    run.underline = True


def add_kop(text):
    global nr
    global doc
    nr += 1
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.style = doc.styles["specs"]
    p.alignment = 0
    run = p.add_run(str(nr) + "." + text)
    run.bold = True
    run.underline = True


def add_foto(names, oneperson=False):
    global doc
    if oneperson:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        names = names.lower()
        klas = klas_e.get()
        try:
            p.add_run().add_picture(f"pictures/{klas}/{names}.jpg", width=Cm(3), height=Cm(4))
        except FileNotFoundError:
            error = Label(status, text=f"Error, {names} zit niet in je klas!")
            error.pack()
        else:
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name = doc.add_paragraph(names)
            name.alignment = 1
            name.paragraph_format.space_after = Pt(0)

    else:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        klas = klas_e.get()
        for name in names:
            name = name.lower()
            try:
                p.add_run().add_picture(f"pictures/{klas}/{name}.jpg", width=Cm(3), height=Cm(4))
            except FileNotFoundError:
                error = Label(status, text=f"Error, {name} zit niet in je klas")
                error.pack()
            else:
                p.add_run(" " * 7)
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2 = doc.add_paragraph()
        p2.alignment = 1
        p2.paragraph_format.space_after = Pt(0)
        for i in range(len(names)):
            p2.add_run(names[i])
            p2.add_run(" " * 7)
            i += 1


def add_info(vraag, info):
    global doc
    p = doc.add_paragraph()
    p.add_run(vraag).bold = True
    p.add_run(info)


def add_lid():
    global leden_aantl
    leden_aantl += 1
    if leden_aantl == 2:
        lid2_e.place(relx=0.25, rely=0.4, relwidth=0.7)
    elif leden_aantl == 3:
        lid3_e.place(relx=0.25, rely=0.6, relwidth=0.7)
    elif leden_aantl == 4:
        lid4_e.place(relx=0.25, rely=0.8, relwidth=0.7)
    elif leden_aantl > 4:
        leden_aantl = 4


def del_lid():
    global leden_aantl
    leden_aantl -= 1
    if leden_aantl == 3:
        lid4_e.place_forget()
    elif leden_aantl == 2:
        lid3_e.place_forget()
    elif leden_aantl == 1:
        lid2_e.place_forget()
    elif leden_aantl < 1:
        leden_aantl = 1


def go_back():
    global leden_aantl

    back.place_forget()
    vak_i.place_forget()
    prac_nr_i.place_forget()
    prac_titel_i.place_forget()
    datum_i.place_forget()
    klas_i.place_forget()
    jaar_i.place_forget()
    ber_i.place_forget()

    naam_i.place_forget()
    leden1_i.place_forget()
    leden2_i.place_forget()

    filename_i.place_forget()
    create_i.place_forget()

    helper.place(relx=0.8, rely=0, relwidth=0.2, relheight=0.8)
    vak_e.place(relx=0.25, rely=0, relwidth=0.7)
    prac_nr_e.place(relx=0.25, rely=0.14, relwidth=0.7)
    prac_titel_e.place(relx=0.25, rely=0.28, relwidth=0.7)
    datum_e.place(relx=0.25, rely=0.42, relwidth=0.7)
    klas_e.place(relx=0.25, rely=0.56, relwidth=0.7)
    jaar_e.place(relx=0.25, rely=0.70, relwidth=0.7)
    ber.place(relx=0.25, rely=0.84)
    naam_e.place(relx=0.25, rely=0, relwidth=0.7)
    lid1_e.place(relx=0.25, rely=0.2, relwidth=0.7)

    if leden_aantl == 2:
        lid2_e.place(relx=0.25, rely=0.4, relwidth=0.7)
    elif leden_aantl == 3:
        lid2_e.place(relx=0.25, rely=0.4, relwidth=0.7)
        lid3_e.place(relx=0.25, rely=0.6, relwidth=0.7)
    elif leden_aantl == 4:
        lid2_e.place(relx=0.25, rely=0.4, relwidth=0.7)
        lid3_e.place(relx=0.25, rely=0.6, relwidth=0.7)
        lid4_e.place(relx=0.25, rely=0.8, relwidth=0.7)

    filename_e.place(relx=0.25, rely=0, relwidth=0.45)
    path_b.place(relx=0.75, rely=0, relwidth=0.2)
    voorb_b.place(relx=0.25, rely=0.5, relwidth=0.2)
    create_file.place(relx=0.55, rely=0.5, relwidth=0.2)


def show_help():
    helper.place_forget()
    vak_e.place_forget()
    prac_nr_e.place_forget()
    prac_titel_e.place_forget()
    datum_e.place_forget()
    klas_e.place_forget()
    jaar_e.place_forget()
    ber.place_forget()
    create_file.place_forget()
    voorb_b.place_forget()

    naam_e.place_forget()
    lid1_e.place_forget()
    if leden_aantl == 2:
        lid2_e.place_forget()
    elif leden_aantl == 3:
        lid2_e.place_forget()
        lid3_e.place_forget()
    elif leden_aantl == 4:
        lid2_e.place_forget()
        lid3_e.place_forget()
        lid4_e.place_forget()

    filename_e.place_forget()
    path_b.place_forget()

    back.place(relx=0.8, rely=0, relwidth=0.2, relheight=0.8)
    vak_i.place(relx=0.25, rely=0, relwidth=0.7)
    prac_nr_i.place(relx=0.25, rely=0.14, relwidth=0.7)
    prac_titel_i.place(relx=0.25, rely=0.28, relwidth=0.7)
    datum_i.place(relx=0.25, rely=0.42, relwidth=0.7)
    klas_i.place(relx=0.25, rely=0.56, relwidth=0.7)
    jaar_i.place(relx=0.25, rely=0.70, relwidth=0.7)
    ber_i.place(relx=0.25, rely=0.84, relwidth=0.7)

    naam_i.place(relx=0.25, rely=0, relwidth=0.7)
    leden1_i.place(relx=0.25, rely=0.3, relwidth=0.7)
    leden2_i.place(relx=0.25, rely=0.5, relwidth=0.7)

    filename_i.place(relx=0.25, rely=0, relwidth=0.7)
    create_i.place(relx=0.25, rely=0.5, relwidth=0.7)


def get_path():
    global path
    path = filedialog.askdirectory()


def get_voorb():
    global voorb_f
    voorb_f = askopenfilename(initialdir="C:/Users/User/Documents/",
                              filetypes=(("Documents", "*.docx"), ("All Files", "*.*")),
                              title="Kies een bestand."
                              )


def find_dvdp(path2):
    doc2 = Document(path2)
    i = 0
    for paragraph in doc2.paragraphs:
        i += 1
        if "Doel van de proef" in paragraph.text:
            break
    dvdp = doc2.paragraphs[i].text
    return dvdp


def check_info():
    run = False
    vak = vak_e.get()
    prac_nr = prac_nr_e.get()
    prac_titel = prac_titel_e.get()
    datum = datum_e.get()
    klas = klas_e.get()
    jaar = jaar_e.get()
    filename = filename_e.get() + ".docx"

    try:
        prac_nr = int(prac_nr)
    except ValueError:
        error = Label(status, text=f"Error, {prac_nr} is geen geldig practicum nummer")
        error.pack()

    naam = naam_e.get()
    namen = naam
    if leden_aantl == 1:
        namen = (naam, lid1_e.get())
    elif leden_aantl == 2:
        namen = (naam, lid1_e.get(), lid2_e.get())
    elif leden_aantl == 3:
        namen = (naam, lid1_e.get(), lid2_e.get(), lid3_e.get())
    elif leden_aantl == 4:
        namen = (naam, lid1_e.get(), lid2_e.get(), lid3_e.get(), lid4_e.get())

    klas_dir = Path(f"pictures/{klas}")
    if klas_dir.is_dir():
        for naam in namen:
            naam = naam.lower() + ".jpg"
            file = Path(f"pictures/{klas}/{naam}")
            if file.is_file():
                run = True
            else:
                error = Label(status, text=f"Error, {naam} zit niet in je klas")
                error.pack()
                break
    else:
        error = Label(status, text=f"Error, {klas} zit nog niet in het systeem!")
        error.pack()

    if run:
        datum_regex = re.compile(r"\d\d/\d\d/\d\d\d\d")
        jaar_regex = re.compile(r"\d\d\d\d - \d\d\d\d")
        if voorb_f is not "niets":
            voorb_path = Path(voorb_f)
            if voorb_path.exists():
                if type(vak) == str:
                    if type(prac_titel) == str:
                        mo = datum_regex.search(datum)
                        if mo is not None:
                            klas_dir = Path(f"pictures/{klas}")
                            if klas_dir.is_dir():
                                mo = jaar_regex.search(jaar)
                                if mo is not None:
                                    file = Path(f"{filename}")
                                    if file.exists():
                                        error = Label(status, text=f"Error, {filename} is al een bestand")
                                        error.pack()
                                    else:
                                        create_doc()
                                else:
                                    error = Label(status, text=f"Error, {jaar} is geen geldig jaar")
                                    error.pack()
                            else:
                                error = Label(status, text=f"Error, {klas} zit nog niet in het systeem")
                                error.pack()
                        else:
                            error = Label(status, text=f"Error, {datum} is geen geldige datum")
                            error.pack()
                    else:
                        error = Label(status, text=f"Error, {prac_titel} is geen geldige Titel")
                        error.pack()
                else:
                    error = Label(status, text=f"Error, {vak} is geen vak!")
                    error.pack()
            else:
                error = Label(status, text=f"Error, De voorbereiding bestaat niet!")
                error.pack()
        else:
            error = Label(status, text=f"Error, Er is geen voobereiding geselecteerd!")
            error.pack()


def create_doc():
    bussy = Label(status, text="creating document")
    bussy.pack()
    global nr
    global doc
    global path

    vak = vak_e.get()
    prac_nr = str(prac_nr_e.get()) + " "
    prac_titel = prac_titel_e.get()
    datum = datum_e.get()
    klas = klas_e.get()
    jaar = jaar_e.get()
    filename = filename_e.get() + ".docx"

    naam = naam_e.get()
    namen = naam
    if leden_aantl == 1:
        namen = (naam, lid1_e.get())
    elif leden_aantl == 2:
        namen = (naam, lid1_e.get(), lid2_e.get())
    elif leden_aantl == 3:
        namen = (naam, lid1_e.get(), lid2_e.get(), lid3_e.get())
    elif leden_aantl == 4:
        namen = (naam, lid1_e.get(), lid2_e.get(), lid3_e.get(), lid4_e.get())

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture("pictures/logo_vti.jpg", width=Cm(5.66), height=Cm(3.59))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_titel("practicum " + str(prac_nr) + vak)
    add_titel(prac_titel)
    add_spec("Het verslag van:")
    add_foto(namen[0], True)
    add_spec("De groepsleden:")
    add_foto(namen)
    add_info("Datum van het practicum: ", datum)
    add_info("Klas: ", klas)
    add_info("Vak: ", vak)
    add_info("Schooljaar: ", jaar)
    add_info("Leerkracht: ", "Ing. B. Aernoudt")
    doc.add_page_break()
    add_kop("Doel van de proef")
    doc.add_paragraph(find_dvdp(voorb_f))
    add_kop("Benodigdheden")
    add_kop("H- en P- zinnen en gevarensymbolen")
    add_kop("Beschrijving, voorstelling en waarnemingen van de proef")
    add_kop("Vragen")
    global berekeningen
    if berekeningen.get() == 1:
        add_kop("Berekeningen")
    add_kop("Besluit")
    doc.save(f"{path}/{filename}")
    bussy.destroy()
    done = Label(status, text="Document klaar!")
    done.pack()
    root.after(7000)
    done.destroy()
    doc.close()


# --------- GUI ---------

HEIGHT = 600
WIDTH = 700

root = Tk()
root.title("Practicum Maker")

berekeningen = IntVar()

canvas = Canvas(root, width=WIDTH, height=HEIGHT)
canvas.pack()

welcome_frame = Frame(root, bd=5)
welcome_frame.place(relx=0.1, rely=0, relwidth=0.8, relheight=0.1)

welcome = Label(welcome_frame, text="Welkom bij Practicum Maker. Vul de gegevens in om een practicum te maken")
welcome.place(relwidth=0.75, relheight=1)

helper = Button(welcome_frame, text="Help", command=show_help)
back = Button(welcome_frame, text="back", command=go_back)
helper.place(relx=0.8, rely=0, relwidth=0.2, relheight=0.8)

# ----------   PRAC INFO   ----------

info_frame = LabelFrame(root, text="Practicum Info", bd=5)
info_frame.place(relx=0.1, rely=0.1, relwidth=0.8, relheight=0.3)

vak_l = Label(info_frame, text="Vak:")
vak_l.place(relx=0.02, rely=0, relwidth=0.2)
vak_e = Entry(info_frame)
vak_e.place(relx=0.25, rely=0, relwidth=0.7)

prac_nr_l = Label(info_frame, text="Practicum Nummer:")
prac_nr_l.place(relx=0.02, rely=0.14, relwidth=0.2)
prac_nr_e = Entry(info_frame)
prac_nr_e.place(relx=0.25, rely=0.14, relwidth=0.7)

prac_titel_l = Label(info_frame, text="Practicum Titel:")
prac_titel_l.place(relx=0.02, rely=0.28, relwidth=0.2)
prac_titel_e = Entry(info_frame)
prac_titel_e.place(relx=0.25, rely=0.28, relwidth=0.7)

datum_l = Label(info_frame, text="Datum:")
datum_l.place(relx=0.02, rely=0.42, relwidth=0.2)
datum_e = Entry(info_frame)
datum_e.place(relx=0.25, rely=0.42, relwidth=0.7)
datum_e.insert(0, "dd/mm/jaar")

klas_l = Label(info_frame, text="Klas:")
klas_l.place(relx=0.02, rely=0.56, relwidth=0.2)
klas_e = Entry(info_frame)
klas_e.place(relx=0.25, rely=0.56, relwidth=0.7)
klas_e.insert(0, "4IW")

jaar_l = Label(info_frame, text="Schooljaar:")
jaar_l.place(relx=0.02, rely=0.70, relwidth=0.2)
jaar_e = Entry(info_frame)
jaar_e.place(relx=0.25, rely=0.70, relwidth=0.7)
jaar_e.insert(0, "2019 - 2020")

ber = Checkbutton(info_frame, text="Berekeningen", variable=berekeningen)
ber.place(relx=0.25, rely=0.84)

# ----------   GROEP   ----------
leden_aantl = 1

groep_frame = LabelFrame(root, text="Groep", bd=5)
groep_frame.place(relx=0.1, rely=0.4, relwidth=0.8, relheight=0.3, )

naam_l = Label(groep_frame, text="Eigen naam:")
naam_l.place(relx=0.02, rely=0, relwidth=0.2)
naam_e = Entry(groep_frame)
naam_e.place(relx=0.25, rely=0, relwidth=0.7)

leden_l = Label(groep_frame, text="Groepsleden:")
leden_l.place(relx=0.02, rely=0.2, relwidth=0.2)

lid1_e = Entry(groep_frame)
lid1_e.place(relx=0.25, rely=0.2, relwidth=0.7)

lid2_e = Entry(groep_frame)
lid3_e = Entry(groep_frame)
lid4_e = Entry(groep_frame)

add = Button(groep_frame, text="+", command=add_lid)
add.place(relx=0.02, rely=0.8, relwidth=0.1)
sub = Button(groep_frame, text="-", command=del_lid)
sub.place(relx=0.13, rely=0.8, relwidth=0.1)

# ----------   BESTAND   ---------

file_frame = LabelFrame(root, text="Bestand", bd=5)
file_frame.place(relx=0.1, rely=0.7, relwidth=0.8, relheight=0.13)

filename_l = Label(file_frame, text="Bestandsnaam:")
filename_l.place(relx=0.02, rely=0, relwidth=0.2)
filename_e = Entry(file_frame)
filename_e.place(relx=0.25, rely=0, relwidth=0.45)

path_b = Button(file_frame, text="locatie", command=get_path)
path_b.place(relx=0.75, rely=0, relwidth=0.2)

voorb_b = Button(file_frame, text="voorbereiding", command=get_voorb)
voorb_b.place(relx=0.25, rely=0.5, relwidth=0.2)

create_file = Button(file_frame, text="Maak Practicum", command=check_info)
create_file.place(relx=0.55, rely=0.5, relwidth=0.2)

# ----------   status   ----------

status = LabelFrame(root, text="status", bd=5)
status.place(relx=0.1, rely=0.83, relwidth=0.8)

# ----------   HELP   ----------

vak_i = Label(info_frame, text="Vul hier het vak in")

prac_nr_i = Label(info_frame, text="Vul het nummer in van het practicum")

prac_titel_i = Label(info_frame, text="Vul het titel van het practicum in")

datum_i = Label(info_frame, text="Datum waarop het practicum werd uitgevoerd")

klas_i = Label(info_frame, text="Vul je klas in (zorg dat deze in het systeem zit)")

jaar_i = Label(info_frame, text="Vul het schooljaar in")

ber_i = Label(info_frame, text="Moet er een kop berekenigen in het verslag staan")

naam_i = Label(groep_frame, text="Vul je eigen naam in (Naam + Voornaam)")

leden1_i = Label(groep_frame, text="Vul de namen in van je groepsleden (Naam + Voornaam)")
leden2_i = Label(groep_frame, text="Druk op + en - om leden toe te voegen of te verwijderen")

filename_i = Label(file_frame, text="Vul de bestandsnaam in en selecteer een bestandlocatie")

create_i = Label(file_frame, text="Selecteer de practicum voorberijding en klik op maak practicum")

root.mainloop()
